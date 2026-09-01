#!/usr/bin/env python3
"""
Content extractors for various document formats.
Extract plain text from PDFs, Word documents, Markdown, and text files.
"""

import os
import re
from pathlib import Path
from typing import Optional


# Column detection: a two-column page has a vertical whitespace band near the
# middle that most rows respect. Anything narrower than MIN_GUTTER is word
# spacing, not a gutter. The band is a fraction of the extent the page's WORDS
# occupy, never of any box declared on the page, so bleed, trim and a shifted
# origin all measure the same.
GUTTER_SEARCH_BAND = (0.38, 0.62)
GUTTER_PROBE_STEP = 0.5
MIN_GUTTER = 8.0
MIN_COLUMN_WORDS = 40
MIN_COLUMN_SHARE = 0.15
# Titles, footers and wide tables cross both columns. They are a minority of the
# rows on a two-column page; past this share the page is simply full-width.
MAX_SPANNING_ROWS = 0.25
# A share measured over a handful of rows is noise: an art page carrying a title,
# a caption and a page number is two thirds "spanning" and still has columns
# under it. Below this many rows the share is not evidence, and the checks that
# follow it — enough words, a wide enough gap, both sides populated — are what
# decide. Above it, a page that is mostly full-width rows is full-width.
MIN_ROWS_TO_VETO = 8
GAP_OUTLIER_ROWS = 0.05
# Trimming votes off a handful of rows manufactures a gap out of nothing, so a
# page has to have this many rows before any are discarded as outliers.
MIN_ROWS_TO_TRIM = 8
# Two boxes overlapping by this much of the shorter one are the same printed
# line.
ROW_OVERLAP_SHARE = 0.8
# Boxes that touch are measured at 0.000 to -0.001pt, so the "gap" at a seam can
# come out very slightly negative. That is float noise in the layout, not two
# words sharing ink; anything past it overlaps for real.
SEAM_MIN = -0.01
# How far apart two word boxes may sit and still be halves of one word. A seam
# opens where the extractor split a word at a change of font, so the two halves
# TOUCH: every one in the sample book measures 0.000 to 0.001pt. A real space is
# an order of magnitude wider — the narrowest in that book is 1.30pt — and the
# margin between the two is what this threshold sits in.
#
# Scaling it to glyph height is tempting and wrong: this book's 18pt display
# type spaces its words 1.45pt apart while its 9pt body type spaces them from
# 1.30pt, so proportion runs backwards, and any ratio wide enough to matter
# welds real words ("TheGrimhammerbrothersare").
SEAM_MAX = 0.5


def robust_gap(rows, probe: float, bounds):
    """The whitespace column around `probe` that all but a few rows agree on.

    Each row votes with its nearest word edge on either side of `probe`. A
    centred title or a stray wide word would pinch the gap shut on its own, so a
    few of the most intrusive votes are dropped from each side. The share is
    deliberately small: trimming hard would eat the real column edge that most
    rows agree on and push the split into the text.
    """
    lo, hi = bounds
    lefts, rights = [], []
    for _, _, words in rows:
        left = [w['x1'] for w in words if w['x1'] <= probe]
        right = [w['x0'] for w in words if w['x0'] >= probe]
        lefts.append(max(left) if left else lo)
        rights.append(min(right) if right else hi)

    lefts.sort()
    rights.sort()
    # On a page of a few rows, discarding even one vote is discarding a fifth of
    # the evidence, and the "gap" left behind can be pure invention that splits a
    # full-width page mid-sentence. Trim only where there is enough to trim.
    drop = max(1, int(len(rows) * GAP_OUTLIER_ROWS)) if len(rows) >= MIN_ROWS_TO_TRIM else 0
    start = lefts[max(0, len(lefts) - 1 - drop)]
    end = rights[min(len(rights) - 1, drop)]
    return start, end


def band_bounds(bbox, band):
    """Absolute (low, high) of a fractional band within a page bbox axis.

    Takes `(start, end)` of the bbox axis rather than a width, because a page
    whose mediabox does not start at the origin has coordinates offset by it —
    fractions of the width alone would point outside the page.
    """
    start, end = bbox
    return start + (end - start) * band[0], start + (end - start) * band[1]


def gutter_from_rows(rows, lo: float, hi: float) -> Optional[float]:
    """Locate the gutter shared by most rows, ignoring the ones that cross it.

    A chapter title or a wide table bridges the two columns. Measuring the page
    as one block lets any single such row veto the split and silently interleave
    the whole page, so the probe first finds the x that the fewest rows cross,
    drops those rows, and measures the gap in what is left.
    """
    if not rows:
        return None

    def crossings(x):
        return sum(1 for _, _, ws in rows if any(w['x0'] < x < w['x1'] for w in ws))

    steps = int((hi - lo) / GUTTER_PROBE_STEP) + 1
    candidates = [lo + i * GUTTER_PROBE_STEP for i in range(steps)]
    center = (lo + hi) / 2
    probe = min(candidates, key=lambda x: (crossings(x), abs(x - center)))
    if len(rows) >= MIN_ROWS_TO_VETO and crossings(probe) > MAX_SPANNING_ROWS * len(rows):
        return None

    kept_rows = [r for r in rows if not any(w['x0'] < probe < w['x1'] for w in r[2])]
    kept = [w for _, _, ws in kept_rows for w in ws]
    if len(kept) < MIN_COLUMN_WORDS:
        return None

    start, end = robust_gap(kept_rows, probe, (lo, hi))
    if end - start < MIN_GUTTER:
        return None

    # The gap above was measured with the most intrusive votes discarded, so its
    # midpoint can land inside a word that was trimmed away — and a kept row with
    # a word across the split is read into BOTH columns, which cuts a paragraph
    # in half around the whole of the other column. The split is therefore pulled
    # back into the untrimmed clearing around the probe, which by construction no
    # kept row crosses.
    split = min(max((start + end) / 2,
                    max((w['x1'] for w in kept if w['x1'] <= probe), default=probe)),
                min((w['x0'] for w in kept if w['x0'] >= probe), default=probe))
    left = sum(1 for w in kept if w['x1'] <= split)
    right = sum(1 for w in kept if w['x0'] >= split)
    if min(left, right) < MIN_COLUMN_SHARE * len(kept):
        return None
    return split


def text_extent(rows):
    """The (x0, x1) the page's words actually occupy."""
    edges = [(w['x0'], w['x1']) for _, _, ws in rows for w in ws]
    return min(x0 for x0, _ in edges), max(x1 for _, x1 in edges)


def find_column_gutter(page, rows=None) -> Optional[float]:
    """X coordinate to split a two-column page at, or None if it is one column.

    The gutter is searched for in the middle of the TEXT, not the middle of the
    paper. Neither box on a page can be trusted to locate that: pdfplumber's
    `Page.bbox` is the mediabox, which on a print-ready PDF carries bleed and
    trim well beyond the printed area, while a cropbox can disagree with the
    content outright on a page whose origin has been shifted. Both aim the
    search off centre, it finds nothing, and the page silently reads as one
    column. The words themselves are the one honest measure, and they need no
    special case for either.

    Cover art and near-empty pages have too few words to judge, and pages whose
    "gutter" leaves one side nearly empty are single-column text with a margin
    note, so both are read whole.
    """
    if rows is None:
        rows = text_rows(page.extract_words())
    if sum(len(ws) for _, _, ws in rows) < MIN_COLUMN_WORDS:
        return None

    lo, hi = band_bounds(text_extent(rows), GUTTER_SEARCH_BAND)
    return gutter_from_rows(rows, lo, hi)


def _same_line(a, b) -> bool:
    """True if two word boxes sit on the same printed line.

    The test is containment, not mere overlap: the shorter box must lie almost
    entirely within the taller one. A display capital contains the small
    capitals beside it, and a drop cap contains the body line it opens, so both
    read as one line. Two lines merely staggered against each other — a heading
    in one column starting partway down the line beside it — overlap by roughly
    three quarters and must stay apart, or the columns weld together.
    """
    overlap = min(a['bottom'], b['bottom']) - max(a['top'], b['top'])
    shortest = min(a['bottom'] - a['top'], b['bottom'] - b['top'])
    return shortest > 0 and overlap >= ROW_OVERLAP_SHARE * shortest


def text_rows(words):
    """Group words into visual rows, top to bottom.

    Returns `[(top, bottom, words)]`. Rows are found by vertical OVERLAP, not by
    a shared top edge: a chapter heading sets its first letter as a display
    capital twice the height of the small capitals beside it, and the two have
    tops several points apart while plainly sitting on the same line. Grouping
    by top edge alone splits that heading in half, and half a heading is how
    "Part 1: The Call" reaches the spine as "Part 1: t".

    A word joins the open row only if it sits on the same line as EVERY word
    already in it. Comparing against the row's accumulated extent lets a drop
    cap three body lines tall rake all three in; comparing against just the last
    arrival is no better, because an intermediate-height word welds two printed
    lines that do not overlap each other at all. Membership has to be pairwise
    with the whole row, and then neither can happen.
    """
    rows = []
    for word in sorted(words, key=lambda w: (w['top'], w['x0'])):
        if rows and all(_same_line(other, word) for other in rows[-1][2]):
            rows[-1][0] = min(rows[-1][0], word['top'])
            rows[-1][1] = max(rows[-1][1], word['bottom'])
            rows[-1][2].append(word)
        else:
            rows.append([word['top'], word['bottom'], [word]])
    return [(top, bottom, sorted(ws, key=lambda w: w['x0'])) for top, bottom, ws in rows]


def row_text(words) -> str:
    """Join one row's words, closing the seam where a word was split mid-glyph.

    A display capital is a separate word object from the letters that follow it
    ("P" then "art"), and they touch. Anything further apart is a real space
    between real words.
    """
    parts = []
    previous = None
    for word in sorted(words, key=lambda w: w['x0']):
        if previous is not None and not words_are_glued(previous, word):
            parts.append(' ')
        parts.append(word['text'])
        previous = word
    return ''.join(parts)


def words_are_glued(left, right) -> bool:
    """True if two boxes are halves of one word rather than two words.

    They have to touch. Anything a space wide is a space, and anything that
    OVERLAPS is not glued either: two words cannot occupy the same ink, so a
    negative gap means the row is out of reading order, and joining there would
    invent a token ("Whisperscity") that appears in no printed line. A row
    grouped wrongly therefore degrades to spaced words, never to a new one.

    The overlap test is signed, not absolute. Italic and tightly kerned type
    comes out of pdfplumber with boxes that overlap by a fraction of a point,
    and reading a fraction of a point of overlap as "touching" welds two real
    words ("meleecombat"). Only the float noise at a true seam is forgiven.
    """
    gap = right['x0'] - left['x1']
    return SEAM_MIN <= gap <= SEAM_MAX


def page_text_by_rows(page, gutter: Optional[float] = None, rows=None) -> str:
    """Read a page in reading order, row by row.

    Rows decide where text goes, not geometry. On a two-column page a row
    crossing the gutter — a chapter title, a boxed read-aloud, a wide table — is
    emitted once, whole; every other run of rows is emitted left column, then
    right column. Without a gutter the page is one run of rows.

    This deliberately does not crop bands out of the page. `crop` keeps any
    glyph that merely overlaps its box, so a band boundary landing inside a row
    emitted that row twice: once as debris in each column, once in the spanning
    band. Walking rows cannot double-emit, because each row is written exactly
    where it is assigned.
    """
    if rows is None:
        rows = text_rows(page.extract_words())
    if gutter is None:
        return '\n'.join(row_text(ws) for _, _, ws in rows)

    blocks = []
    for spans_gutter, run in _row_runs(rows, gutter):
        if spans_gutter:
            blocks.append('\n'.join(row_text(ws) for _, _, ws in run))
            continue
        columns = [[], []]
        for _, _, ws in run:
            for column, side in zip(columns, (
                [w for w in ws if w['x1'] <= gutter],
                [w for w in ws if w['x0'] >= gutter],
            )):
                if side:
                    column.append(row_text(side))
        blocks.extend('\n'.join(c) for c in columns if c)
    return '\n\n'.join(b for b in blocks if b.strip())


def _row_runs(rows, gutter: float):
    """Consecutive rows batched by whether they cross the gutter."""
    runs = []
    for row in rows:
        spans_gutter = any(w['x0'] < gutter < w['x1'] for w in row[2])
        if runs and runs[-1][0] == spans_gutter:
            runs[-1][1].append(row)
        else:
            runs.append((spans_gutter, [row]))
    return runs


class PDFExtractor:
    """Extract text content from PDF files."""

    def __init__(self):
        """Initialize PDF extractor."""
        self.pypdf_available = False
        self.pdfplumber_available = False
        # Which reader last produced text. A caller that asked for columns and
        # silently got 'plain' has interleaved text in its hands, and needs to
        # be able to say so rather than ship it as if it were clean.
        self.last_mode = None

        try:
            import pypdf
            self.pypdf_available = True
            self.pypdf = pypdf
        except ImportError:
            pass

        try:
            import pdfplumber
            self.pdfplumber_available = True
            self.pdfplumber = pdfplumber
        except ImportError:
            pass

        if not self.pypdf_available and not self.pdfplumber_available:
            print("Warning: No PDF libraries available. Install pypdf or pdfplumber.")

    def extract(self, filepath: str, column_aware: bool = False) -> str:
        """
        Extract text from a PDF file.

        Args:
            filepath: Path to the PDF file
            column_aware: Detect two-column pages and read each column in turn
                instead of letting the extractor interleave them line by line.
                Requires pdfplumber; ignored (with a warning) without it.

        Returns:
            Extracted text content
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"PDF file not found: {filepath}")

        self.last_mode = None
        if column_aware:
            if self.pdfplumber_available:
                try:
                    text, degraded = self._extract_columns_with_pdfplumber(filepath)
                    # Every page falling through to a plain read leaves the
                    # caller holding exactly the interleaved text the flag was
                    # meant to avoid, so it is reported as the plain read it is.
                    self.last_mode = 'plain' if degraded else 'columns'
                    return text
                except Exception as e:
                    # A PDF pdfplumber cannot open should come back as text
                    # rather than a traceback out of the slicer — but that text
                    # is single-column, which is the very interleaving the flag
                    # exists to avoid, so `last_mode` reports what really ran.
                    print(f"Column-aware extraction failed: {e}")
                    print("Falling back to single-column extraction; columns will interleave.")
            else:
                print("Warning: column-aware extraction needs pdfplumber; using default mode.")

        # Try pdfplumber first (better for complex layouts)
        if self.pdfplumber_available:
            try:
                text = self._extract_with_pdfplumber(filepath)
                self.last_mode = 'plain'
                return text
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")
                if self.pypdf_available:
                    print("Falling back to pypdf...")

        # Fall back to pypdf
        if self.pypdf_available:
            try:
                text = self._extract_with_pypdf(filepath)
                self.last_mode = 'plain'
                return text
            except Exception as e:
                print(f"pypdf extraction failed: {e}")

        raise RuntimeError("No PDF extraction library available. Install pypdf or pdfplumber.")

    def _extract_with_pdfplumber(self, filepath: str) -> str:
        """Extract text using pdfplumber."""
        text = []
        with self.pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"--- Page {page_num} ---\n")
                        text.append(page_text)
                        text.append("\n\n")
                except Exception as e:
                    print(f"Error extracting page {page_num}: {e}")

        return ''.join(text)

    def _extract_columns_with_pdfplumber(self, filepath: str):
        """Extract text page by page, splitting two-column pages at the gutter.

        Adventure modules are laid out in two columns; a straight extract reads
        across the page and shuffles the two columns together sentence by
        sentence. Here each page is measured for a vertical whitespace band near
        its middle, and when one is found the page is read column by column.

        Returns `(text, degraded)`. A page the row reader chokes on falls back to
        a plain read one page at a time, and enough of those is a book that was
        not read in columns at all — which the caller has to be able to say.
        """
        text = []
        pages_read = fell_back = 0
        with self.pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                pages_read += 1
                try:
                    rows = text_rows(page.extract_words())
                    gutter = find_column_gutter(page, rows)
                    page_text = page_text_by_rows(page, gutter, rows)
                except Exception as e:
                    # Losing a page costs its text AND shifts the page numbers
                    # every later scene cites, so fall back to a plain read
                    # rather than dropping it.
                    print(f"Error reading columns on page {page_num}: {e}")
                    fell_back += 1
                    try:
                        page_text = page.extract_text()
                    except Exception as inner:
                        print(f"Error extracting page {page_num}: {inner}")
                        page_text = ''
                if page_text:
                    text.append(f"--- Page {page_num} ---\n")
                    text.append(page_text)
                    text.append("\n\n")

        return ''.join(text), bool(fell_back) and fell_back * 2 >= pages_read

    def _extract_with_pypdf(self, filepath: str) -> str:
        """Extract text using pypdf."""
        text = []
        with open(filepath, 'rb') as file:
            pdf_reader = self.pypdf.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"--- Page {page_num + 1} ---\n")
                        text.append(page_text)
                        text.append("\n\n")
                except Exception as e:
                    print(f"Error extracting page {page_num + 1}: {e}")

        return ''.join(text)


class MarkdownExtractor:
    """Extract text content from Markdown files."""

    def extract(self, filepath: str) -> str:
        """
        Extract text from a Markdown file.

        Args:
            filepath: Path to the Markdown file

        Returns:
            Text content with basic Markdown formatting preserved
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Markdown file not found: {filepath}")

        with open(filepath, 'r', encoding='utf-8') as file:
            content = file.read()

        # Clean up excessive whitespace while preserving structure
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content


class DocxExtractor:
    """Extract text content from Word documents."""

    def __init__(self):
        """Initialize DOCX extractor."""
        self.docx_available = False

        try:
            import docx
            self.docx_available = True
            self.docx = docx
        except ImportError:
            print("Warning: python-docx not available. Install python-docx to process Word documents.")

    def extract(self, filepath: str) -> str:
        """
        Extract text from a Word document.

        Args:
            filepath: Path to the DOCX file

        Returns:
            Extracted text content
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Word document not found: {filepath}")

        if not self.docx_available:
            # Try basic text extraction as fallback
            return self._basic_extract(filepath)

        try:
            doc = self.docx.Document(filepath)
            text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
                    text.append('\n')

            # Extract tables
            for table in doc.tables:
                text.append('\n--- Table ---\n')
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text.append(' | '.join(row_text))
                    text.append('\n')
                text.append('\n')

            return ''.join(text)

        except Exception as e:
            print(f"Error extracting Word document: {e}")
            return self._basic_extract(filepath)

    def _basic_extract(self, filepath: str) -> str:
        """Basic text extraction fallback for Word documents."""
        # This is a very basic fallback that may not work well
        try:
            with open(filepath, 'rb') as file:
                content = file.read()
                # Try to decode as UTF-8, ignoring errors
                text = content.decode('utf-8', errors='ignore')
                # Remove binary garbage
                text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
                return text
        except Exception as e:
            raise RuntimeError(f"Cannot extract text from Word document: {e}")


class TextExtractor:
    """Extract content from plain text files."""

    def extract(self, filepath: str) -> str:
        """
        Extract text from a plain text file.

        Args:
            filepath: Path to the text file

        Returns:
            File content
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Text file not found: {filepath}")

        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue

        # If all encodings fail, try with error handling
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()


def extract_content(filepath: str) -> str:
    """
    Extract text content from any supported file type.

    Args:
        filepath: Path to the file

    Returns:
        Extracted text content
    """
    file_path = Path(filepath)
    extension = file_path.suffix.lower()

    extractors = {
        '.pdf': PDFExtractor(),
        '.md': MarkdownExtractor(),
        '.markdown': MarkdownExtractor(),
        '.docx': DocxExtractor(),
        '.doc': DocxExtractor(),
        '.txt': TextExtractor(),
        '.text': TextExtractor(),
    }

    extractor = extractors.get(extension)
    if not extractor:
        raise ValueError(f"Unsupported file type: {extension}")

    return extractor.extract(filepath)


class ContentExtractor:
    """Unified content extractor for any supported file type."""

    def extract_text(self, filepath: str) -> str:
        """Extract text from any supported file type."""
        return extract_content(filepath)


def main():
    """Test the extractors."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: content_extractor.py <file>")
        sys.exit(1)

    filepath = sys.argv[1]

    try:
        content = extract_content(filepath)
        print(f"Extracted {len(content)} characters from {filepath}")
        print("\nFirst 500 characters:")
        print("-" * 50)
        print(content[:500])
        print("-" * 50)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()